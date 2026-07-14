"""Builds the two-page SeekhAI pitch note as an editable Word .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x2F, 0x5D, 0x50)
DARK = RGBColor(0x14, 0x20, 0x1B)
GREY = RGBColor(0x55, 0x60, 0x5A)

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.06

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.75)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = GREEN
    return p


def para(text, after=5, size=10.5, color=DARK, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color
    return p


def bullet(label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if label:
        r = p.add_run(label + "  "); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = DARK
    return p


# ---- Title -----------------------------------------------------------------
t = doc.add_paragraph()
t.paragraph_format.space_after = Pt(1)
r = t.add_run("SeekhAI"); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = GREEN
para("Interactive, adaptive pre-class learning — a smarter alternative to AI-generated videos.",
     after=8, size=12, color=GREY, italic=True)

# ---- Opening ---------------------------------------------------------------
para("Colleges are turning to AI-generated video to prepare students before class. But "
     "producing video is expensive, slow to update, and the narration is often subtly "
     "inaccurate — and students still watch passively and arrive under-prepared. SeekhAI "
     "takes a different path. Instead of a video to watch, it turns your own course material "
     "into an interactive, adaptive course that students learn by doing.")

# ---- From your material ----------------------------------------------------
h1("From your material to a live course — in minutes")
bullet("Upload what you already have.", "Lecture slides, PDFs, notes and readings. SeekhAI "
       "reads them and builds a structured course — topics, subtopics and a full set of "
       "interactions — grounded first and foremost in your content.")
bullet("Fills the gaps from the live web.", "Where a concept needs more, SeekhAI researches "
       "current, reputable sources (papers, official docs) to enrich and keep material "
       "up to date — while your material stays the source of truth.")
bullet("No talking-head videos.", "Every idea arrives as something the student actively does.")

# ---- Why it beats video ----------------------------------------------------
h1("Why it beats uninteractive video")
bullet("Active, not passive.", "Multiple-choice, order-the-steps, fill-in-the-blanks, "
       "drag-and-drop diagrams, and guided worked-example / code walkthroughs.")
bullet("Adaptive difficulty.", "Each topic is graded (Difficulty Level 1 to 3). The course "
       "steps up as a student succeeds and eases back when they struggle.")
bullet("Dynamic routing.", "A wrong answer routes the student to a targeted follow-up that "
       "pinpoints the exact gap; a streak of correct answers unlocks harder material. No "
       "student is left stuck, and none is left bored.")
bullet("Accurate by design.", "Content is grounded in your material and independently "
       "checked before it reaches a student. You stay in control of what is taught.")
bullet("Lower, predictable cost.", "No per-minute rendering and no re-shoots. Updating a "
       "course means re-uploading a file — not regenerating a video.")

# ---- Teacher visibility ----------------------------------------------------
h1("Daily visibility for teachers — over WhatsApp")
bullet("A daily digest to each teacher.", "Who engaged, how far they got, where the class is "
       "struggling, and which topics to reinforce in the next session — delivered simply over "
       "WhatsApp, no new app to learn.")
bullet("Pre-class prep becomes a live signal.", "Teachers walk in knowing exactly where "
       "students are, instead of guessing — turning preparation from a black box into data.")

# ---- Expandable ------------------------------------------------------------
h1("Built to expand across your entire catalogue")
para("SeekhAI is discipline-agnostic. It builds rich, interactive courses from any "
     "department's material — and is designed to grow well beyond computer science:", after=4)
bullet("Core engineering.", "Auto-generated, interactive simulations — circuits, mechanics, "
       "signals, thermodynamics — so students manipulate systems, not just read about them.")
bullet("Management & MBA.", "Case-style decision interactions, scenario branching and "
       "data-driven exercises that build judgement, not recall.")
bullet("Sciences, humanities & professional programmes.", "The same learn-by-doing engine, "
       "seeded from each department's own material.")
para("One platform that grows with you — far more expandable than a library of static videos "
     "that must be re-produced every time a syllabus changes.", after=6)

# ---- Pilot -----------------------------------------------------------------
h1("A simple, low-risk pilot")
bullet("Pick one course.", "A high-enrolment, pre-class-heavy subject. Hand us the existing "
       "material; we build the interactive course and wire up the daily teacher digest.")
bullet("Run it with one cohort.", "For a few weeks, measured against clear goals: higher "
       "pre-class completion, better in-class readiness, and teacher time saved.")
bullet("Decide from evidence.", "Expand to more courses and departments only once the pilot "
       "proves the model on your own students.")

# ---- Comparison table ------------------------------------------------------
h1("At a glance: AI video vs. SeekhAI")
rows = [
    ("", "AI-generated video", "SeekhAI"),
    ("Student experience", "Passive — watch and forget", "Active — learn by doing"),
    ("Accuracy", "Often subtly wrong; hard to fix", "Grounded in your material + verified"),
    ("Adapts to the student", "No — one video for everyone", "Yes — difficulty & routing per student"),
    ("Teacher insight", "None", "Daily WhatsApp engagement + performance"),
    ("Updating content", "Re-generate / re-shoot", "Re-upload the file"),
    ("Expandability", "New videos per topic", "One engine across every discipline"),
]
table = doc.add_table(rows=len(rows), cols=3)
table.style = "Light Grid Accent 1"
table.autofit = True
for ci, val in enumerate(rows[0]):
    c = table.rows[0].cells[ci]
    c.text = ""
    rr = c.paragraphs[0].add_run(val); rr.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = GREEN
for ri in range(1, len(rows)):
    for ci in range(3):
        c = table.rows[ri].cells[ci]
        c.text = ""
        rr = c.paragraphs[0].add_run(rows[ri][ci])
        rr.font.size = Pt(9.5)
        rr.bold = (ci == 0)
        rr.font.color.rgb = DARK

# ---- Close -----------------------------------------------------------------
para("")
close = doc.add_paragraph()
close.paragraph_format.space_before = Pt(2)
r = close.add_run("The takeaway: ")
r.bold = True; r.font.color.rgb = GREEN; r.font.size = Pt(10.5)
r2 = close.add_run("for the price and effort of one AI-video library, SeekhAI gives your "
                   "students an interactive, adaptive, always-current course built from your "
                   "own material — and gives your teachers a daily line of sight into how the "
                   "class is doing. We would love to prove it on one of your courses.")
r2.font.size = Pt(10.5); r2.font.color.rgb = DARK

para("Contact: [Your name] · [email] · [phone]", after=0, size=9.5, color=GREY, italic=True)

out = "/Users/Work/seekhoai/pitch/SeekhAI-Pre-Class-Learning.docx"
doc.save(out)
print("saved:", out)
